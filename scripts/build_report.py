from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


TEMPLATE = Path(r"F:\浏览器下载\EDGE下载\综合实践（阶段1）-实验2-实验报告模版.docx")
OUTPUT = Path("reports/综合实践（阶段1）-实验2-实验报告-已完成.docx")


def main() -> None:
    repo_root = Path(__file__).resolve().parents[1]
    doc = Document(TEMPLATE)
    _clear_body(doc)
    _setup_styles(doc)

    _title(doc, "软件产品综合开发实践-（实验二）—实验报告")
    _body(doc, "1. 姓名：__________________    2. 学号：__________________")

    _heading(doc, "3. 请提供你的Agent代码存放的repo（请包含核心的prompt）：")
    _body(doc, "git@github.com:FanFenMo/agent-experiment.git")
    _body(doc, "核心 prompt 位于 prompts/system_prompt.md；Agent 逻辑位于 src/agent_experiment/agent.py；工具定义位于 src/agent_experiment/tools.py。")

    _heading(doc, "4. 请简要介绍你设计的Agent")
    _body(
        doc,
        "我设计的 Agent 名为 Experiment Delivery Agent，是一个离线可运行的课程实验交付助手。"
        "它的主要功能是读取实验二 rubric，将截止时间、交付物、评分项和反思要求拆解成可执行的报告计划，"
        "再检查仓库是否包含 prompt、context、工具定义、demo 脚本和测试文件。Agent 使用的工具包括 read_context、"
        "extract_rubric、plan_report、draft_reflection 和 validate_repository。上下文集成方式包括文件型 rubric context、"
        "系统 prompt、工具 schema，以及 JSON execution trace。项目不依赖外部 API，因此老师可以在本地稳定复现实验运行过程。"
    )

    _heading(doc, "5. 请提供你的Agent运行时的说明")
    _body(
        doc,
        "运行方式是在仓库根目录设置 PYTHONPATH=src 后执行 scripts/run_demo.py。"
        "demo 会先注册工具，再读取 context/experiment2_rubric.md，随后提取 rubric、规划报告、生成反思草稿并检查仓库完整性。"
        "运行记录保存在 demo_outputs/experiment2_run.json，下面三张截图来自同一次真实执行。"
    )
    _picture(
        doc,
        repo_root / "reports" / "assets" / "execution_1.png",
        "图 1：Agent 启动后加载系统 prompt，并列出 5 个已注册工具。",
    )
    _picture(
        doc,
        repo_root / "reports" / "assets" / "execution_2.png",
        "图 2：Agent 调用 read_context、extract_rubric 和 plan_report，将实验要求转成报告计划。",
    )
    _picture(
        doc,
        repo_root / "reports" / "assets" / "execution_3.png",
        "图 3：Agent 生成反思草稿并检查仓库文件，确认 required files complete。",
    )

    _heading(doc, "6. 简要介绍你在完成实验二过程中如何使用了大语言模型等AI技术帮助你完成实验任务")
    _body(
        doc,
        "我主要把大语言模型用于需求拆解、项目结构设计、代码初稿生成和报告文字整理。实际过程中遇到的具体问题是："
        "AI 有时会把工具参数、执行步骤和文档渲染要求混在一起，输出看起来完整，但如果不落到真实文件和命令上就不可复现。"
        "为了解决这个问题，我先让 AI 查看仓库和模板，再把 Agent 的工具 schema 固定在代码中，要求每一步都输出 JSON execution trace。"
        "随后通过单元测试、demo 运行、截图生成和 DOCX 渲染检查逐项验证。这个过程说明 AI 可以显著提高搭建速度，"
        "但最终仍然需要用明确约束、可运行脚本和可视化验收来保证实验交付质量。"
    )

    output_path = repo_root / OUTPUT
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(output_path)


def _clear_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def _setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)


def _title(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(14)
    run = paragraph.add_run(text)
    _font(run, size=16, bold=True)


def _heading(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    _font(run, size=11.5, bold=True)


def _body(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Inches(0.24)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    _font(run, size=10.5)


def _picture(doc: Document, image_path: Path, caption: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(5.65))

    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_after = Pt(6)
    caption_run = caption_paragraph.add_run(caption)
    _font(caption_run, size=9)


def _font(run, size: float, bold: bool = False) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold


if __name__ == "__main__":
    main()
